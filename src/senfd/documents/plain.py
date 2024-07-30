import re
from pathlib import Path
from typing import ClassVar, List, NamedTuple, Optional, Tuple

import docx
from pydantic import BaseModel, Field

import senfd.errors
import senfd.schemas
import senfd.tables
from senfd.documents.base import Converter, Document, DocumentMeta
from senfd.tables import Cell, Row, Table


class Figure(BaseModel):
    """
    A figure as captioned in a table of figures in the NVMe Specification Documents.

    This is a minimally enriched figure representation. The intent is that via regular
    expressions, it should be possible to construct instances when feeding it with
    the text from a table of figures or a table that conventionally contains a figure
    caption in the first row.

    The page number is only retrievable from table of figure captions and thus optional.
    The table data is similarly represented in a form with minimal enrichment.
    """

    REGEX_TABLE_OF_FIGURES: ClassVar[str] = (
        r"^(?P<caption>Figure\s+(?P<figure_nr>\d+)\s*:"
        r"\s*(?P<description>.*?))(?P<page_nr>\d+)?$"
    )
    REGEX_TABLE_ROW: ClassVar[str] = (
        r"^(?P<caption>Figure\s+(?P<figure_nr>\d+)\s*:" r"\s*(?P<description>.*?))$"
    )
    figure_nr: int  # Figure as numbered in the specification document
    caption: str  # The entire figure caption
    description: str  # The part of figure caption without the "Fig X:" prefix

    page_nr: Optional[int] = None
    table: Optional[senfd.tables.Table] = None

    @classmethod
    def from_regex(cls, regex, text):
        match = re.match(regex, text)
        if not match:
            return None

        data = {
            "figure_nr": int(match.group("figure_nr")),
            "caption": match.group("caption").strip(),
            "description": match.group("description").strip(),
        }
        if "page_nr" in match.groupdict():
            data["page_nr"] = int(match.group("page_nr"))

        return cls(**data)


class FigureDocument(Document):

    SUFFIX_JSON: ClassVar[str] = ".plain.figure.document.json"
    SUFFIX_HTML: ClassVar[str] = ".plain.figure.document.html"

    FILENAME_SCHEMA: ClassVar[str] = "plain.figure.document.schema.json"
    FILENAME_HTML_TEMPLATE: ClassVar[str] = "plain.figure.document.html.jinja2"

    figures: List[Figure] = Field(default_factory=list)


class FromDocx(Converter):

    @staticmethod
    def is_applicable(path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    @staticmethod
    def convert(path: Path) -> Tuple[FigureDocument, List[NamedTuple]]:

        def docx_table_to_table(docx_table: docx.table.Table) -> Table:
            table = Table()

            for docx_row in docx_table.rows:
                row = Row()
                table.rows.append(row)

                for docx_cell in docx_row.cells:
                    cell = Cell(
                        text=str(docx_cell.text),
                        tables=[
                            docx_table_to_table(nested_table)
                            for nested_table in docx_cell.tables
                        ],
                    )
                    row.cells.append(cell)

            return table

        figures = {}
        errors: List[NamedTuple] = []

        docx_document = docx.Document(path)

        # Add tabular figures -- page_nr unavailable
        for table_nr, docx_table in enumerate(docx_document.tables, 1):
            caption = str(docx_table.rows[0].cells[0].text).strip()

            figure = Figure.from_regex(Figure.REGEX_TABLE_ROW, caption)
            if not figure:
                errors.append(
                    senfd.errors.TableCaptionError(
                        "Does not match figure caption assumptions",
                        caption,
                    )
                )
                continue

            if figure.figure_nr in figures:
                errors.append(
                    senfd.errors.TableCaptionError(
                        "Duplicate caption",
                        caption,
                    )
                )
                continue

            figure.table = docx_table_to_table(docx_table)
            figures[figure.figure_nr] = figure

        # Update tabular figures with page_nr
        # Add non-fabular figures
        # Check table-of-figure description validity
        prev = cur = None
        for paragraph in docx_document.paragraphs:
            cur = paragraph.style.name

            # We exit early to avoid scanning the entire document, since we know that
            # once we are looking at a "table of figures" paragraph, then once we see
            # one that is not, then no more will arrive
            if prev == "table of figures" and cur != "table of figures":
                break
            prev = cur
            if paragraph.style.name != "table of figures":
                continue

            # Check whether the paragraph is a reference to a figure
            caption = paragraph.text.strip()
            figure = Figure.from_regex(Figure.REGEX_TABLE_OF_FIGURES, caption)
            if not figure:
                errors.append(
                    senfd.errors.TableOfFiguresError(
                        "Does not match figure assumptions",
                        caption,
                    )
                )
                continue

            if not figure.page_nr:
                errors.append(
                    senfd.errors.TableOfFiguresError("Is missing <page_nr>", caption)
                )
                continue

            existing = figures.get(figure.figure_nr, None)
            if existing:
                existing.page_nr = figure.page_nr
                if figure.description not in existing.description:
                    errors.append(
                        senfd.errors.TableOfFiguresError(
                            f"({existing.description}) != {figure.description}",
                            caption,
                        )
                    )
            else:
                figures[figure.figure_nr] = figure

        return (
            FigureDocument(
                meta=DocumentMeta(stem=path.stem), figures=list(figures.values())
            ),
            errors,
        )
